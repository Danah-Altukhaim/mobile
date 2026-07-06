#!/usr/bin/env python3
"""Dump every CCK Student Hub doc (.docx/.xlsx) to plain text for inspection.
Not part of the data build — a convenience for reading the source docs.
"""
import os, sys, glob
import docx
import openpyxl

REPO = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
DOCS = os.path.join(REPO, 'CCK Student Hub Docs')
OUT = os.path.join(os.path.dirname(__file__), '.cache', 'extracted')
os.makedirs(OUT, exist_ok=True)

def extract_docx(path):
    d = docx.Document(path)
    lines = []
    # iterate body elements in order (paragraphs + tables)
    from docx.oxml.ns import qn
    body = d.element.body
    para_idx = 0
    tbl_idx = 0
    paras = d.paragraphs
    tables = d.tables
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            if para_idx < len(paras):
                t = paras[para_idx].text.strip()
                if t:
                    lines.append(t)
                para_idx += 1
        elif child.tag == qn('w:tbl'):
            if tbl_idx < len(tables):
                tbl = tables[tbl_idx]
                lines.append("\n[TABLE]")
                for row in tbl.rows:
                    cells = [c.text.strip().replace("\n"," / ") for c in row.cells]
                    lines.append(" | ".join(cells))
                lines.append("[/TABLE]\n")
                tbl_idx += 1
    return "\n".join(lines)

def extract_xlsx(path):
    wb = openpyxl.load_workbook(path, data_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"\n===== SHEET: {ws.title} (dims {ws.dimensions}) =====")
        rows = list(ws.iter_rows(values_only=True))
        for r in rows:
            if r is None:
                continue
            vals = ["" if v is None else str(v).strip() for v in r]
            if any(vals):
                out.append(" | ".join(vals))
    return "\n".join(out)

for path in sorted(glob.glob(DOCS + "/*")):
    base = os.path.basename(path)
    if base.startswith("~$") or base.startswith("."):
        continue
    ext = base.lower().rsplit(".",1)[-1]
    try:
        if ext == "docx":
            content = extract_docx(path)
        elif ext == "xlsx":
            content = extract_xlsx(path)
        else:
            continue
    except Exception as e:
        content = f"ERROR extracting: {e}"
    outname = base.rsplit(".",1)[0].replace("/","_") + ".txt"
    with open(os.path.join(OUT, outname), "w") as f:
        f.write(f"# SOURCE: {base}\n\n{content}\n")
    print(f"{base}: {len(content)} chars -> {outname}")
